#!/usr/bin/env python3
"""
rehydrate.py – Pseudonym-Rehydrierung für Moodle-Agent-MCP

Ersetzt Pseudonyme (S-0001 …) in .docx/.md/.txt-Dateien durch Klarnamen.
Die Klarnamen kommen aus pseudonym-map.json (NIEMALS an Modell/Cloud weitergeben!).

CLI-Verwendung:
  python rehydrate.py --infile bericht.md --outfile bericht_klar.md --field fullname
  python rehydrate.py --infile ergebnisse.docx           # in-place

Mit --json-result (für MCP-Server-Aufruf):
  Gibt JSON aus: {"ok": true, "datei": "...", "pseudonyme_ersetzt": N,
                  "schueler": ["S-0001", ...], "field": "..."}
  KEINE Klarnamen in der JSON-Ausgabe.

Benötigt für .docx:  pip install python-docx
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path


# ---------------------------------------------------------------------------
# Gemeinsame Kernfunktionen (importierbar von anderen Skripten)
# ---------------------------------------------------------------------------

def load_map(map_path: str) -> dict:
    """Lädt pseudonym-map.json und gibt das users-Dict zurück."""
    if not os.path.exists(map_path):
        raise FileNotFoundError(
            f"pseudonym-map.json nicht gefunden: {map_path}"
        )
    with open(map_path, encoding="utf-8") as f:
        data = json.load(f)
    return data.get("users", {})


def build_replacements(
    users: dict, field: str
) -> list:
    """
    Baut [(Regex, Ersetzung, Pseudonym), …]-Liste auf.
    Längste Pseudonyme zuerst, um Teilüberlappungen zu vermeiden.
    """
    pairs = []
    for _uid, info in users.items():
        ps = info.get("pseudonym", "")
        if not ps:
            continue
        if field == "email":
            target = info.get("email", "")
        elif field == "username":
            target = info.get("username", "")
        else:  # fullname (Standard)
            fn = info.get("fullname", "")
            target = (
                fn
                or info.get("firstname", "")
                or (fn.split()[0] if fn else "")
            )
        if not target:
            continue
        pattern = re.compile(
            r"(?<![^\W_])" + re.escape(ps) + r"(?![^\W_])",
            re.IGNORECASE | re.UNICODE,
        )
        pairs.append((pattern, target, ps))
    # Längste Pseudonyme (= längste Muster) zuerst
    pairs.sort(key=lambda x: len(x[0].pattern), reverse=True)
    return pairs


def apply_replacements(text: str, pairs: list) -> tuple:
    """
    Ersetzt alle Pseudonyme.
    Gibt (neuer_text, Menge_ersetzter_Pseudonyme) zurück.
    """
    replaced = set()
    for pattern, target, ps in pairs:
        new_text, n = pattern.subn(target, text)
        if n > 0:
            replaced.add(ps)
            text = new_text
    return text, replaced


# ---------------------------------------------------------------------------
# .txt / .md
# ---------------------------------------------------------------------------

def rehydrate_text_file(infile: str, outfile: str, pairs: list) -> set:
    with open(infile, encoding="utf-8") as f:
        content = f.read()
    new_content, replaced = apply_replacements(content, pairs)
    with open(outfile, "w", encoding="utf-8") as f:
        f.write(new_content)
    return replaced


# ---------------------------------------------------------------------------
# .docx  (python-docx)
# ---------------------------------------------------------------------------

def _rehydrate_paragraph(para, pairs: list, replaced: set) -> None:
    """
    Führt Runs zusammen, ersetzt Pseudonyme und schreibt den Text zurück.
    Der erste Run behält die Formatierung; restliche Runs werden geleert.
    """
    full_text = "".join(r.text for r in para.runs)
    new_text, new_replaced = apply_replacements(full_text, pairs)
    if new_replaced:
        replaced.update(new_replaced)
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]:
                r.text = ""


def rehydrate_docx_file(infile: str, outfile: str, pairs: list) -> set:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        raise ImportError(
            "python-docx nicht installiert. Bitte: pip install python-docx"
        )

    doc = Document(infile)
    replaced: set = set()

    # Normaler Fließtext
    for para in doc.paragraphs:
        _rehydrate_paragraph(para, pairs, replaced)

    # Tabellen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _rehydrate_paragraph(para, pairs, replaced)

    # Kopf- und Fußzeilen
    for section in doc.sections:
        for region in (section.header, section.footer):
            if region is not None:
                for para in region.paragraphs:
                    _rehydrate_paragraph(para, pairs, replaced)

    doc.save(outfile)
    return replaced


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def _auto_map_path() -> str:
    """Sucht pseudonym-map.json relativ zum Skript."""
    return str(Path(__file__).parent / "pseudonym-map.json")


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Rehydriert Pseudonyme (S-xxxx) in Berichtsdateien."
    )
    parser.add_argument("--infile", required=True,
                        help="Eingabedatei (.docx/.md/.txt)")
    parser.add_argument("--outfile", default=None,
                        help="Ausgabedatei (Standard: in-place, überschreibt infile)")
    parser.add_argument(
        "--field", default="fullname",
        choices=["fullname", "email", "username"],
        help="Welches Klardaten-Feld einsetzen (Standard: fullname)",
    )
    parser.add_argument("--map-path", default=None,
                        help="Pfad zur pseudonym-map.json (Standard: neben dem Skript)")
    parser.add_argument(
        "--json-result", action="store_true",
        help="Ergebnis als JSON ausgeben (für MCP-Server-Aufruf, keine Klarnamen)",
    )
    args = parser.parse_args()

    outfile = args.outfile or args.infile
    map_path = args.map_path or _auto_map_path()

    try:
        users = load_map(map_path)
        pairs = build_replacements(users, args.field)

        ext = Path(args.infile).suffix.lower()
        if ext == ".docx":
            replaced = rehydrate_docx_file(args.infile, outfile, pairs)
        elif ext in (".md", ".txt"):
            replaced = rehydrate_text_file(args.infile, outfile, pairs)
        else:
            raise ValueError(f"Nicht unterstützter Dateityp: {ext!r}")

        result = {
            "ok": True,
            "datei": Path(outfile).name,
            "pseudonyme_ersetzt": len(replaced),
            "schueler": sorted(replaced),
            "field": args.field,
        }

    except Exception as exc:  # noqa: BLE001
        result = {
            "ok": False,
            "datei": Path(outfile).name if outfile else "",
            "pseudonyme_ersetzt": 0,
            "schueler": [],
            "field": args.field,
            "fehler": str(exc),
        }
        if not args.json_result:
            print(f"Fehler: {exc}", file=sys.stderr)
            sys.exit(1)

    if args.json_result:
        print(json.dumps(result, ensure_ascii=False))
    else:
        if result["ok"]:
            print(
                f"✓ {result['pseudonyme_ersetzt']} Pseudonym(e) ersetzt → {outfile}"
            )
        else:
            print(f"✗ Fehler: {result.get('fehler', 'unbekannt')}", file=sys.stderr)
            sys.exit(1)


if __name__ == "__main__":
    main()
